"""
Word document loader.

Supports:
- .docx (Office Open XML format)
- .doc (legacy format, requires additional libraries)
"""

import logging
from pathlib import Path
from typing import Optional

from ..base_loader import BaseLoader, LoaderException

logger = logging.getLogger(__name__)


class WordLoader(BaseLoader):
    """
    Load and extract text from Word documents.

    Supports .docx files using python-docx library.
    """

    SUPPORTED_EXTENSIONS = ['.docx', '.doc']

    def load(self) -> str:
        """
        Load text from Word document.

        Returns:
            Extracted text content
        """
        path = Path(self.source)
        self._validate_file_exists(path)

        if path.suffix == '.docx':
            return self._load_docx(path)
        elif path.suffix == '.doc':
            return self._load_doc_legacy(path)
        else:
            raise LoaderException(f"Unsupported Word format: {path.suffix}")

    def _load_docx(self, path: Path) -> str:
        """Extract text from .docx file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        logger.info(f"Loading Word document: {path}")

        try:
            doc = Document(path)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)

            # Combine all text
            content = "\n\n".join(paragraphs)

            if table_texts:
                content += "\n\n" + "\n".join(table_texts)

            # Store metadata
            self.metadata = {
                'file_path': str(path),
                'file_size': path.stat().st_size,
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'char_count': len(content),
                'format': 'docx'
            }

            logger.info(
                f"Extracted {len(paragraphs)} paragraphs, "
                f"{len(doc.tables)} tables, {len(content)} chars"
            )

            return content

        except Exception as e:
            raise LoaderException(f"Error loading Word document: {e}")

    def _load_doc_legacy(self, path: Path) -> str:
        """
        Extract text from legacy .doc file.

        Note: This is challenging and may require external tools.
        For now, we raise an error and suggest converting to .docx.
        """
        raise LoaderException(
            f"Legacy .doc format not directly supported. "
            f"Please convert {path.name} to .docx format first. "
            f"You can use LibreOffice or Microsoft Word to convert: "
            f"Open the file and Save As -> Word Document (.docx)"
        )
